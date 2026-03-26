#!/usr/bin/env python3
"""Generate Pamir AI market research Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(20, 60, 120)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_kv_table(doc, items):
    """Add a key-value table."""
    table = doc.add_table(rows=len(items), cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(items):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if table.rows[i].cells[0].paragraphs[0].runs else None
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
    return table

def add_data_table(doc, headers, rows):
    """Add a data table with headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)
    return table

def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============================================================
# Cover
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Pamir AI')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(20, 60, 120)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('市场调研报告')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(60, 60, 60)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年3月  |  Edge AI Agent Hardware')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ============================================================
# 1. Company Overview
# ============================================================
doc.add_heading('1. 公司概况', level=1)

add_kv_table(doc, [
    ('公司名称', 'PamirAI Inc.'),
    ('成立时间', '2024年'),
    ('总部', 'San Francisco, CA, USA'),
    ('创始人', 'Kevin (Chengming) Zhang, Tianqi Ye'),
    ('定位', 'AI agents on affordable hardware, distributing personal intelligence'),
    ('官网', 'https://www.pamir.ai/'),
    ('阶段', 'Early-stage startup (Open Alpha)'),
])

doc.add_paragraph()
doc.add_paragraph(
    'Pamir AI 是一家 2024 年成立于旧金山的 AI 硬件初创公司，致力于将 AI Agent '
    '运行在低成本的专用硬件上，让每个人都能拥有属于自己的 AI「大脑」。公司核心理念类似于 '
    'Raspberry Pi 对计算的普及——让 AI 开发不再依赖昂贵的订阅和高性能云服务器。'
)

# ============================================================
# 2. Funding
# ============================================================
doc.add_heading('2. 融资情况', level=1)

add_kv_table(doc, [
    ('融资阶段', 'Pre-seed / Early-stage'),
    ('估值', '未公开'),
])

doc.add_paragraph()
doc.add_heading('已知投资方', level=2)
add_bullets(doc, [
    'Founders, Inc.',
    'Samsung NEXT Ventures',
    'iSeed Ventures',
    'Llama Ventures',
    'Neo',
])

doc.add_paragraph(
    'Pamir AI 尚未有公开的正式定价融资轮次，但已获得多家知名机构支持。'
    '其中 Samsung NEXT 的参与尤为值得关注，意味着该公司在硬件供应链方面可能获得战略资源。'
)

# ============================================================
# 3. Products
# ============================================================
doc.add_heading('3. 产品线', level=1)

doc.add_heading('3.1 Distiller Alpha Dev Kit ($250)', level=2)
doc.add_paragraph(
    '口袋大小的 Linux 计算机，专为 AI Agent 24/7 无人值守运行设计。'
    '预装 ClawdBot (Claude Code agent)，扫码即可在浏览器中打开 VS Code 远程开发。'
    'Alpha 批次已售罄，日产 10-15 台，交付周期 2-3 周。'
)

doc.add_heading('硬件规格', level=3)
add_kv_table(doc, [
    ('处理器', '64-bit ARM @ 2.4GHz'),
    ('内存', '8GB RAM'),
    ('存储', '32GB eMMC'),
    ('连接', 'Wi-Fi, Bluetooth, USB-C'),
    ('I/O 协议', 'I2C/SPI, UART, GPIO, PWM'),
    ('显示', 'E-ink 显示屏'),
    ('音频', '内置麦克风 + 扬声器'),
    ('本地 AI', '3B 参数模型'),
])

doc.add_paragraph()
doc.add_heading('本地推理性能', level=3)
add_data_table(doc,
    ['模型', '参数量', '速度'],
    [
        ['Qwen3', '0.6B', '~21 tokens/s'],
        ['Qwen3', '1.7B', '~9 tokens/s'],
    ]
)

doc.add_paragraph()
doc.add_heading('软件生态', level=3)
add_bullets(doc, [
    'OS: Custom Linux (Distiller Agent Layer)',
    '预装: ClawdBot (Claude Code agent)',
    '远程开发: 浏览器内 VS Code，扫码即用',
    '移动端: iOS App (Beta)，桌面端即将推出',
    'SDK: DistillerSDK (GitHub 开源)',
])

doc.add_heading('3.2 Agent Computer — 下一代 (2026 夏季)', level=2)
doc.add_paragraph(
    '从零开始设计的专用 Agent 计算机。搭载 8 核 ARM 处理器、内置电池、全新外形设计。'
    'Alpha 批次已售罄，目前开放预约注册。将是 Pamir 的关键里程碑产品。'
)

# ============================================================
# 4. Target Market
# ============================================================
doc.add_heading('4. 目标市场', level=1)

doc.add_heading('核心用户群', level=2)
add_bullets(doc, [
    '开发者 — 需要 24/7 AI 编码助手、自动测试和部署',
    '硬件工程师 — 嵌入式系统开发 (刷机、测试、迭代)',
    '创客/Maker — 非专业人士想用硬件构建项目',
    '隐私敏感用户 — 希望本地运行 AI，数据不离开设备',
])

doc.add_heading('核心使用场景', level=2)
add_bullets(doc, [
    'Agentic Workflows: AI Agent 自主编码、测试、部署，持续数小时/天',
    '嵌入式开发: 通过 USB-C 多协议接口直接与传感器和执行器交互',
    '远程开发: 随时随地通过浏览器扫码使用 VS Code',
    'IoT 原型: 智能硬件快速原型开发',
])

# ============================================================
# 5. Competitive Landscape
# ============================================================
doc.add_heading('5. 竞争格局', level=1)

doc.add_heading('产品对比', level=2)
add_data_table(doc,
    ['产品', '价格', 'AI 算力', '优势', '劣势'],
    [
        ['Pamir Distiller', '$250', '3B 本地', 'Agent 专用，开箱即用', '产能有限'],
        ['RPi 5 + AI HAT+', '$80-150', '13-26 TOPS', '生态庞大，社区活跃', '需自行配置软件'],
        ['Jetson Orin Nano', '$249-499', '40-275 TOPS', 'GPU 高性能推理', '功耗高、价格贵'],
        ['Google Coral', '$60-150', '4 TOPS', 'Google 生态', '性能有限、更新慢'],
        ['BeagleY-AI', '$70-100', '4 TOPS', '集成 AI 加速', '生态不成熟'],
    ]
)

doc.add_paragraph()
doc.add_heading('Pamir AI 的差异化优势', level=2)
add_bullets(doc, [
    '市场上唯一专为 AI Agent 24/7 运行设计的硬件产品',
    '预装 ClawdBot (Claude Code)，开箱即用',
    'E-ink 显示屏 + 低功耗 ARM，适合持续运行',
    'USB-C 统一接口支持多种硬件协议 (I2C/SPI/UART/GPIO/PWM)',
    '扫码即可在浏览器中使用 VS Code 远程开发',
    '软硬件一体化方案，大幅降低入门门槛',
])

# ============================================================
# 6. Market Context
# ============================================================
doc.add_heading('6. 市场背景', level=1)

add_kv_table(doc, [
    ('2025 边缘 AI 市场规模', '$26.14B'),
    ('2030 预测规模', '$58.90B'),
    ('年复合增长率 (CAGR)', '17.6%'),
])

doc.add_paragraph()
doc.add_heading('关键趋势', level=2)
add_bullets(doc, [
    'IoT 设备爆发式增长推动边缘 AI 需求',
    '隐私和数据主权意识增强，本地推理受青睐',
    'AI Agent 从云端走向端侧，24/7 自主运行成为新趋势',
    '开发者工具与硬件深度整合 (Claude Code + 专用硬件)',
    '亚太地区增长最快，制造业与半导体产业驱动',
    'Raspberry Pi 6 可能内置 NPU，推动 AI 在单板计算机上的普及',
])

doc.add_heading('主要玩家', level=2)
doc.add_paragraph('Qualcomm, Huawei, Samsung, Apple, MediaTek, NVIDIA, Google, IBM')

# ============================================================
# 7. SWOT
# ============================================================
doc.add_heading('7. SWOT 分析', level=1)

doc.add_heading('Strengths (优势)', level=2)
add_bullets(doc, [
    '独特定位: 市场唯一专为 AI Agent 设计的口袋计算机',
    '价格优势: $250 远低于 Mac Mini ($600)',
    '软硬件一体: 预装 Claude Code，开箱即用',
    '强力投资方: Samsung NEXT 等知名机构背书',
    '开发者友好: 扫码 VS Code、iOS App、开源 SDK',
])

doc.add_heading('Weaknesses (劣势)', level=2)
add_bullets(doc, [
    '产能有限: 每天仅 10-15 台，难以规模化',
    '本地模型能力受限: 3B 参数模型智能程度有限',
    '早期阶段: Alpha 产品，稳定性和生态待验证',
    '依赖第三方 AI: 核心 Agent 能力依赖 Anthropic Claude API',
    '团队规模小: 早期创业团队，资源有限',
])

doc.add_heading('Opportunities (机会)', level=2)
add_bullets(doc, [
    'AI Agent 市场爆发: 2026 年 Agent 应用进入主流',
    '边缘 AI 市场快速增长 (CAGR 17.6%)',
    '开发者对专用 AI 硬件需求旺盛 (Alpha 秒售罄)',
    '夏季推出下一代产品，有望大幅提升竞争力',
    '可拓展到教育、IoT、智能家居等垂直领域',
])

doc.add_heading('Threats (威胁)', level=2)
add_bullets(doc, [
    'Raspberry Pi 6 可能内置 NPU，蚕食低端市场',
    'NVIDIA、Qualcomm 等巨头可能推出类似定位产品',
    'AI 推理能力快速提升，专用硬件可能被通用设备取代',
    '开源社区可在树莓派上复刻类似软件方案',
    '供应链风险: 小批量生产成本高、交付不确定',
])

# ============================================================
# 8. Assessment
# ============================================================
doc.add_heading('8. 综合评估', level=1)

add_kv_table(doc, [
    ('市场契合度', '高 — AI Agent 从概念走向实际应用，开发者需要专用硬件'),
    ('风险等级', '中高 — 早期产品，竞争激烈，需快速迭代和规模化'),
])

doc.add_paragraph()
doc.add_paragraph(
    'Pamir AI 在 AI Agent 专用硬件这一新兴细分市场中占据先发优势。$250 的定价和软硬件一体化'
    '方案形成了独特卖点。Alpha 批次秒售罄说明市场需求真实存在。但产能瓶颈、本地模型能力有限、'
    '以及来自树莓派生态和芯片巨头的潜在竞争是主要挑战。2026 年夏季推出的下一代 Agent Computer '
    '将是关键里程碑。'
)

doc.add_heading('重点关注事项', level=2)
add_bullets(doc, [
    '2026 年夏季 Agent Computer 的发布和市场反应',
    '是否进行正式融资轮 (Seed / Series A)',
    '产能提升和供应链建设进展',
    '与 Anthropic / Claude 的合作深度',
    '开发者社区增长和生态建设',
])

# ============================================================
# Sources
# ============================================================
doc.add_heading('参考来源', level=1)
sources = [
    'https://www.pamir.ai/',
    'https://shop.pamir.ai/products/distiller-two-alpha-dev-kit',
    'https://f.inc/portfolio/pamir/',
    'https://pitchbook.com/profiles/company/756474-40',
    'https://tracxn.com/d/companies/pamir/',
    'https://aiagentstore.ai/ai-agent/pamir-ai',
    'https://513.toys/distiller-alpha/',
    'https://pamir-ai.hashnode.dev/',
    'https://docs.pamir.ai/development',
    'https://www.marketsandmarkets.com/Market-Reports/edge-ai-hardware-market-158498281.html',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# Save
out_path = '/home/user/daily-ai-pulse-data/data/pamir-market-research.docx'
doc.save(out_path)
print(f'Word document generated: {out_path}')
